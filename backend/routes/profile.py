import json
import os
import re
import tempfile
import unicodedata

from flask import Blueprint, jsonify, request
from werkzeug.utils import secure_filename

from db import get_db
from services.career_ai_service import call_llm
from services.ml_recommender import split_skills


profile_bp = Blueprint("profile", __name__, url_prefix="/api/profile")

ALLOWED_EXTENSIONS = {"pdf", "doc", "docx", "txt", "md"}
REQUIRED_SOFT_ABILITIES = {
    "创新能力": {"score": 3, "description": "能主动尝试新方法解决问题。"},
    "学习能力": {"score": 4, "description": "能持续学习岗位相关知识。"},
    "抗压能力": {"score": 3, "description": "能在项目周期内稳定推进任务。"},
    "沟通能力": {"score": 3, "description": "能清晰表达并协作完成任务。"},
    "实习能力": {"score": 3, "description": "具备基础实践意识，可通过项目继续提升。"},
}


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def _json_loads(value, default):
    if not value:
        return default
    try:
        return json.loads(value)
    except Exception:
        return default


def ensure_required_soft_abilities(soft_abilities):
    result = dict(REQUIRED_SOFT_ABILITIES)
    if isinstance(soft_abilities, dict):
        result.update({k: v for k, v in soft_abilities.items() if v})
    return result


def _rule_extract_profile(text, data=None):
    data = data or {}
    skills = split_skills(data.get("skills_certs") or data.get("skills") or "")
    skill_pool = [
        "Python", "Java", "SQL", "Vue", "React", "Flask", "Django", "Docker",
        "ECharts", "Pandas", "NumPy", "Excel", "Tableau", "Power BI",
        "机器学习", "深度学习", "数据分析", "数据清洗", "可视化", "产品设计",
        "软件测试", "单元测试", "测试用例", "缺陷跟踪", "JMeter", "Postman",
        "测试工程师",
    ]
    skills += [s for s in skill_pool if s.lower() in text.lower()]
    skills = list(dict.fromkeys(skills))
    certs = [item for item in split_skills(data.get("skills_certs") or "") if "证" in item or "PMP" in item.upper()]
    education = data.get("education", "")
    work = data.get("work", "")
    project = data.get("project", "")
    sections = _split_resume_sections(text)
    if not education:
        education = sections.get("education", "")
    if not work:
        work = sections.get("work", "")
    if not project:
        project = sections.get("project", "")
    return {
        "skills": skills,
        "certificates": certs,
        "soft_abilities": ensure_required_soft_abilities({}),
        "education": _parse_education(f"{education}\n{text}"),
        "work_experience": _parse_list_text(work, ["company", "position", "description"]),
        "project_experience": _parse_list_text(project, ["project_name", "role", "description"]),
    }


def _parse_education(text):
    if not text:
        return {}
    degree = ""
    for item in ["博士", "硕士", "本科", "大专", "专科"]:
        if item in text:
            degree = item
            break
    major = ""
    for item in ["软件工程", "计算机科学与技术", "计算机", "数据科学", "人工智能", "信息管理", "电子商务", "市场营销"]:
        if item in text:
            major = item
            break
    grade = ""
    for item in ["大一", "大二", "大三", "大四", "研一", "研二", "研三", "应届"]:
        if item in text:
            grade = item
            break
    return {"school": text[:30], "major": major, "degree": degree, "grade": grade}


def _parse_list_text(text, keys):
    if not text:
        return []
    lines = [line.strip(" -·\t") for line in str(text).splitlines() if line.strip()]
    if not lines:
        return []
    items = []
    for line in lines[:4]:
        if "：" in line:
            title, desc = line.split("：", 1)
        elif ":" in line:
            title, desc = line.split(":", 1)
        else:
            title, desc = line[:30], line
        items.append({keys[0]: title[:30], keys[1]: "", keys[2]: desc.strip() or line})
    return items


def _split_resume_sections(text):
    sections = {"education": "", "work": "", "project": "", "skills": ""}
    current = ""
    for raw in str(text or "").splitlines():
        line = raw.strip()
        if not line:
            continue
        normalized = line.replace("：", ":")
        key = ""
        if any(word in line for word in ["教育", "学历", "本科", "硕士", "博士", "专业"]):
            key = "education"
        if any(word in line for word in ["实习", "工作", "经历"]):
            key = "work"
        if any(word in line for word in ["项目", "作品"]):
            key = "project"
        if any(word in line for word in ["技能", "证书"]):
            key = "skills"
        current = key or current
        if current:
            value = normalized.split(":", 1)[1].strip() if ":" in normalized else line
            sections[current] = (sections[current] + "\n" + value).strip()
    return sections


def _extract_with_llm(text, data=None):
    if not text or len(str(text).strip()) < 8:
        return _rule_extract_profile(text, data)
    if not _should_enhance_with_ai():
        return _rule_extract_profile(text, data)
    prompt = f"""
请从以下大学生简历/档案文本中抽取结构化 JSON：
字段：skills, certificates, soft_abilities, education, work_experience, project_experience。
只返回 JSON。

文本：
{text}
"""
    try:
        raw = call_llm(prompt, temperature=0.2, max_tokens=1800)
        start = raw.find("{")
        end = raw.rfind("}")
        if start >= 0 and end > start:
            parsed = json.loads(raw[start:end + 1])
            parsed["soft_abilities"] = ensure_required_soft_abilities(parsed.get("soft_abilities"))
            return parsed
    except Exception:
        pass
    return _rule_extract_profile(text, data)


def _should_enhance_with_ai():
    return request.args.get("ai") == "1" if request else False


def _education_dict(value):
    if isinstance(value, dict):
        return value
    if isinstance(value, list):
        for item in value:
            if isinstance(item, dict):
                return item
    return {}


@profile_bp.route("/submit", methods=["POST"])
def submit_profile():
    data = request.get_json(silent=True) or {}
    user_id = data.get("user_id")
    text = "\n".join(str(data.get(k, "")) for k in ["education", "work", "project", "skills_certs", "summary"])
    ability = _extract_with_llm(text, data)
    education = _education_dict(ability.get("education"))
    major = data.get("major") or education.get("major", "")

    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute(
            """
            INSERT INTO student (
                user_id, name, major, grade, skills, certificates, internships,
                interests, completeness, competitiveness, phone, email,
                education_text, work_text, project_text, skills_certs_text, summary,
                soft_abilities, education_json, work_json, project_json
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                user_id,
                data.get("name"),
                major,
                data.get("grade") or education.get("degree", ""),
                json.dumps(ability.get("skills", []), ensure_ascii=False),
                json.dumps(ability.get("certificates", []), ensure_ascii=False),
                data.get("work", ""),
                json.dumps(data.get("interests", []), ensure_ascii=False),
                80,
                70,
                data.get("phone"),
                data.get("email"),
                data.get("education", ""),
                data.get("work", ""),
                data.get("project", ""),
                data.get("skills_certs", ""),
                data.get("summary", ""),
                json.dumps(ability.get("soft_abilities", {}), ensure_ascii=False),
                json.dumps(education, ensure_ascii=False),
                json.dumps(ability.get("work_experience", []), ensure_ascii=False),
                json.dumps(ability.get("project_experience", []), ensure_ascii=False),
            ),
        )
        conn.commit()
        student_id = cur.lastrowid
    finally:
        conn.close()

    return jsonify({
        "student_id": student_id,
        "skills": ability.get("skills", []),
        "certificates": ability.get("certificates", []),
        "soft_abilities": ability.get("soft_abilities", {}),
    })


@profile_bp.route("/upload", methods=["POST"])
def upload_resume():
    temp_path = ""
    suffix = ""
    try:
        if "file" not in request.files:
            return jsonify({"error": "缺少文件"}), 400
        file = request.files["file"]
        if not file.filename or not allowed_file(file.filename):
            return jsonify({"error": "不支持的文件类型"}), 400
        suffix = file.filename.rsplit(".", 1)[1].lower()
        filename = secure_filename(file.filename) or f"resume.{suffix}"
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
            file.save(tmp.name)
            temp_path = tmp.name
        text, parse_warning = _read_resume_text(temp_path, suffix)
        ability = _extract_with_llm(text)
        return jsonify({
            "text": text,
            "parse_warning": parse_warning,
            "skills": ability.get("skills", []),
            "certificates": ability.get("certificates", []),
            "soft_abilities": ability.get("soft_abilities", {}),
            "education_json": ability.get("education", {}),
            "work_json": ability.get("work_experience", []),
            "project_json": ability.get("project_experience", []),
        })
    except Exception as exc:
        return jsonify({
            "text": "",
            "parse_warning": _friendly_parse_warning(suffix),
            "skills": [],
            "certificates": [],
            "soft_abilities": ensure_required_soft_abilities({}),
            "education_json": {},
            "work_json": [],
            "project_json": [],
        }), 200
    finally:
        try:
            if temp_path:
                os.remove(temp_path)
        except OSError:
            pass


def _read_resume_text(path, suffix):
    if suffix in {"txt", "md"}:
        return open(path, "r", encoding="utf-8", errors="ignore").read(), ""
    if suffix == "pdf":
        errors = []
        text = ""
        for reader in (_read_pdf_with_pypdfium2, _read_pdf_with_pdfplumber, _read_pdf_with_pdfminer):
            try:
                text = reader(path)
                if text.strip():
                    return _clean_resume_text(text), ""
            except Exception as exc:
                errors.append(str(exc))
        return "", _friendly_parse_warning("pdf")
    if suffix in {"doc", "docx"}:
        try:
            from docx import Document
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs), ""
        except Exception as exc:
            return "", _friendly_parse_warning("docx")
    return "", ""


def _friendly_parse_warning(suffix):
    if suffix == "pdf":
        return "PDF 没有抽取到可用文本，可能是扫描件、加密文件或排版过于复杂。请换成 DOCX/TXT，或复制简历内容到手动填写。"
    if suffix in {"doc", "docx"}:
        return "Word 文档没有抽取到可用文本。请另存为 DOCX/TXT，或复制简历内容到手动填写。"
    return "简历解析没有完成。请换成 DOCX/TXT，或复制简历内容到手动填写。"


def _read_pdf_with_pypdfium2(path):
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(path)
    chunks = []
    try:
        for index in range(len(pdf)):
            page = pdf[index]
            try:
                text_page = page.get_textpage()
                try:
                    chunks.append(text_page.get_text_range() or "")
                finally:
                    text_page.close()
            finally:
                page.close()
    finally:
        pdf.close()
    return "\n".join(chunks)


def _read_pdf_with_pdfplumber(path):
    import pdfplumber
    with pdfplumber.open(path) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def _read_pdf_with_pdfminer(path):
    from pdfminer.high_level import extract_text
    return extract_text(path) or ""


def _clean_resume_text(text):
    text = unicodedata.normalize("NFKC", str(text or ""))
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


@profile_bp.route("/<int:student_id>", methods=["GET"])
def get_profile(student_id):
    conn = get_db()
    try:
        row = conn.execute("SELECT * FROM student WHERE id = ?", (student_id,)).fetchone()
    finally:
        conn.close()
    if not row:
        return jsonify({"error": "学生不存在"}), 404
    data = dict(row)
    for key, default in [
        ("skills", []),
        ("certificates", []),
        ("soft_abilities", {}),
        ("education_json", {}),
        ("work_json", []),
        ("project_json", []),
        ("interest_scores", {}),
    ]:
        data[key] = _json_loads(data.get(key), default)
    return jsonify(data)
